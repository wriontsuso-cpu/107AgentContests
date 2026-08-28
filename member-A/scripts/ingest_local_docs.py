#!/usr/bin/env python3
"""从 Downloads 中的科大本地文档提取文本，生成结构化资源。"""

from __future__ import annotations

import json
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from pypdf import PdfReader

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "data" / "local_docs"
TEXT_DIR = OUT_DIR / "extracted_text"
CURATED_OUT = OUT_DIR / "local_doc_resources.json"

SOURCES = [
    {
        "file": Path(r"d:\Downloads\中科大校园网讲解v1.1.2 (2).pdf"),
        "title": "中科大校园网讲解（v1.1.2）",
        "category": "网站入口",
        "category_id": "service_portal",
        "tags": ["网络", "VPN", "邮箱", "ustc-only", "流程", "入口"],
        "content_kind": "流程",
        "how_to": "阅读校园网讲解文档：注册网络、Wi‑Fi、网络通、VPN、邮箱等按文档步骤操作。",
    },
    {
        "file": Path(r"d:\Downloads\科大学生社团画册2025.pdf"),
        "title": "科大学生社团画册 2025",
        "category": "二课/团学活动",
        "category_id": "campus_life_club",
        "tags": ["社团", "二课", "ustc-only", "机会"],
        "content_kind": "机会",
        "how_to": "查阅社团画册了解各社团简介与招新信息，按画册/社团联系方式加入；活动报名可结合青春科大。",
    },
    {
        "file": Path(r"d:\Downloads\大雾实验不完全指北_20240305 (1).pdf"),
        "title": "大雾实验不完全指北",
        "category": "竞赛/科创",
        "category_id": "opportunity_competition",
        "tags": ["实验室", "科研", "学生经验", "ustc-only", "流程"],
        "content_kind": "流程",
        "how_to": "按《大雾实验不完全指北》了解进组、实验安全与日常规范；具体以实验室最新要求为准。",
    },
    {
        "file": Path(r"d:\Downloads\10 关于2026年“极客中心”实验室招募学生开展科研实践的通知.doc"),
        "title": "2026年「极客中心」实验室招募学生开展科研实践",
        "category": "竞赛/科创",
        "category_id": "opportunity_competition",
        "tags": ["实验室", "科研实践", "极客中心", "ustc-only", "机会"],
        "content_kind": "机会",
        "how_to": "阅读招募通知，按通知要求提交申请材料/报名；关注截止时间与联系人。",
    },
    {
        "file": Path(r"d:\Downloads\中国科学技术大学体育课缓修（退课）申请单(1).doc"),
        "title": "体育课缓修（退课）申请单",
        "category": "办事指南",
        "category_id": "service_guide",
        "tags": ["体育课", "退课", "缓修", "表格", "ustc-only", "流程"],
        "content_kind": "流程",
        "how_to": "填写《体育课缓修（退课）申请单》，按教务/体育教学部要求签字盖章后提交。",
    },
    {
        "file": Path(r"d:\Downloads\第五教学楼电教设备投屏操作简介(1).pdf"),
        "title": "第五教学楼电教设备投屏操作简介",
        "category": "办事指南",
        "category_id": "service_guide",
        "tags": ["教室", "投屏", "电教", "ustc-only", "流程"],
        "content_kind": "流程",
        "how_to": "在五教使用电教设备时，按投屏操作简介连接笔记本/手机投屏；故障联系电教支持。",
    },
    {
        "file": Path(r"d:\Downloads\Guide_of_USTC.pdf"),
        "title": "Guide of USTC（校园生活指南）",
        "category": "新生事务",
        "category_id": "campus_life_freshman",
        "tags": ["新生", "校园生活", "ustc-only", "入口"],
        "content_kind": "入口",
        "how_to": "通读 Guide of USTC，按章节查找宿舍、餐饮、交通、办事等生活信息。",
    },
    {
        "file": Path(r"d:\Downloads\USTC培养方案与教务指南 (7).pdf"),
        "title": "USTC 培养方案与教务指南",
        "category": "教务选课",
        "category_id": "academic_teach",
        "tags": ["培养方案", "选课", "教务", "ustc-only", "流程"],
        "content_kind": "流程",
        "how_to": "查阅培养方案与教务指南，核对学分要求与选课规则；具体以教务处当年通知为准。",
    },
    {
        "file": Path(r"d:\Downloads\其他问题回答：(1)(2).docx"),
        "title": "其他问题回答（学生常见问答）",
        "category": "办事指南",
        "category_id": "service_guide",
        "tags": ["问答", "学生经验", "ustc-only", "流程"],
        "content_kind": "流程",
        "how_to": "按文档中的问答条目查找对应问题的处理方式；若与官网冲突，以官方最新通知为准。",
    },
    {
        "file": Path(r"d:\Downloads\2025级本科新生入学须知0828.doc"),
        "title": "2025级本科新生入学须知",
        "category": "新生事务",
        "category_id": "campus_life_freshman",
        "tags": ["新生", "入学", "报到", "ustc-only", "流程"],
        "content_kind": "流程",
        "how_to": "按入学须知完成报到、缴费、宿舍、军训等相关事项；时间节点以当年通知为准。",
    },
]


def extract_pdf(path: Path, max_pages: int = 30) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages[:max_pages]):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


def extract_docx(path: Path) -> str:
    # Prefer python-docx; fall back to raw document.xml if media CRC broken
    if Document is not None:
        try:
            doc = Document(str(path))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception as exc:
            print(f"  python-docx failed ({exc}), fallback to XML")

    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    except Exception:
        return extract_doc_ole(path)
    text = re.sub(r"</w:p>", "\n", xml)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"&amp;", "&", text)
    text = re.sub(r"&lt;", "<", text)
    text = re.sub(r"&gt;", ">", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_doc_ole(path: Path) -> str:
    """粗糙提取老式 .doc 中的可读文本。"""
    raw = path.read_bytes()
    # Prefer UTF-16LE chunks common in Word .doc
    texts: list[str] = []
    try:
        decoded = raw.decode("utf-16-le", errors="ignore")
        cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", decoded)
        cleaned = re.sub(r"\s{2,}", " ", cleaned)
        # Keep CJK-heavy segments
        for m in re.finditer(r"[\u4e00-\u9fffA-Za-z0-9，。；：、（）()\-_/]{8,}", cleaned):
            texts.append(m.group(0))
    except Exception:
        pass
    if len("".join(texts)) < 80:
        # fallback latin-1 printable
        decoded = raw.decode("latin-1", errors="ignore")
        cleaned = re.sub(r"[^\x09\x0a\x0d\x20-\x7e\u4e00-\u9fff]+", " ", decoded)
        texts = re.findall(r"[\u4e00-\u9fff]{2,}[^,]{0,40}", cleaned)
    return "\n".join(texts[:500])


def summarize(text: str, limit: int = 280) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return ""
    return text[:limit]


def build_resource(meta: dict, text: str, local_url: str) -> dict:
    summary = summarize(text) or meta["how_to"]
    return {
        "title": meta["title"],
        "url": local_url,
        "source": "本地文档入库（成员A）",
        "category": meta["category"],
        "published_at": "",
        "summary": summary,
        "content": text[:8000],
        "crawled_at": datetime.now(timezone.utc).isoformat(),
        "tags": meta["tags"],
        "cost": "免费（文档查阅）",
        "how_to": meta["how_to"],
        "relevance_score": 10,
        "kind": "curated",
        "source_type": "local_doc",
        "category_id": meta["category_id"],
        "access_type": "public",
        "content_kind": meta["content_kind"],
        "ustc_specificity": "ustc-only" if "ustc-only" in meta["tags"] else "学生经验",
        "local_file": str(meta["file"]),
        "local_filename": meta["file"].name,
    }


def main() -> int:
    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    # Copy docs into repo data folder for relative links
    docs_store = OUT_DIR / "files"
    docs_store.mkdir(exist_ok=True)

    resources: list[dict] = []
    for meta in SOURCES:
        path: Path = meta["file"]
        if not path.exists():
            print(f"MISSING: {path}")
            continue

        dest = docs_store / path.name
        if not dest.exists() or dest.stat().st_size != path.stat().st_size:
            dest.write_bytes(path.read_bytes())

        suffix = path.suffix.lower()
        print(f"Extracting: {path.name}")
        if suffix == ".pdf":
            # Large pictorial PDFs: fewer pages
            max_pages = 8 if path.stat().st_size > 5_000_000 else 25
            text = extract_pdf(path, max_pages=max_pages)
        elif suffix == ".docx":
            text = extract_docx(path)
        elif suffix == ".doc":
            text = extract_doc_ole(path)
        else:
            text = ""

        text_path = TEXT_DIR / (path.stem[:60] + ".txt")
        text_path.write_text(text or "(无可提取文本，可能为扫描件/图片型 PDF)", encoding="utf-8")

        local_url = f"file:///{dest.resolve().as_posix()}"
        resources.append(build_resource(meta, text, local_url))
        print(f"  chars={len(text)} -> {meta['title']}")

    CURATED_OUT.write_text(json.dumps(resources, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\nWrote {len(resources)} resources -> {CURATED_OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
